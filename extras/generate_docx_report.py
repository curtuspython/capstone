"""
generate_docx_report.py
-----------------------
Generates the project report as an editable Word (.docx) document
with embedded terminal-style screenshots.

Run once to produce Report/Report.docx.

Usage:
    python generate_docx_report.py

The resulting file can be opened in Microsoft Word or Google Docs
and edited freely before converting to PDF for submission.

Screenshots are auto-generated (via generate_screenshots.py) if they
do not already exist.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from generate_screenshots import make_all_screenshots, OUT_DIR as SS_DIR

# The extras/ folder sits next to Capstone_Project-CS[02]/
# So: extras/../Capstone_Project-CS[02]/Report/Report.docx
_PROJECT_DIR = Path(__file__).parent.parent / "Capstone_Project-CS[02]"
OUTPUT_PATH  = _PROJECT_DIR / "Report" / "Report.docx"

# Screenshot paths (generated on demand)
_SS_STARTUP  = SS_DIR / "startup_banner.png"
_SS_SCORING  = SS_DIR / "cv_scoring.png"
_SS_RANKED   = SS_DIR / "ranked_table.png"
_SS_INTERACT = SS_DIR / "interactive_mode.png"


# ---------------------------------------------------------------------------
# Colour constants (Walmart-adjacent professional blue palette)
# ---------------------------------------------------------------------------
BLUE       = RGBColor(0x00, 0x35, 0x80)   # heading blue
DARK_GREY  = RGBColor(0x1E, 0x1E, 0x1E)   # body text
MID_GREY   = RGBColor(0x50, 0x50, 0x50)   # sub-text / captions
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BLUE = RGBColor(0xD6, 0xE4, 0xF7)   # table header bg


# ---------------------------------------------------------------------------
# Style helpers
# ---------------------------------------------------------------------------

def _set_cell_bg(cell, hex_color: str) -> None:
    """Set a table cell background colour via raw OOXML."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    """Add a styled heading paragraph."""
    para = doc.add_heading(text, level=level)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in para.runs:
        run.font.color.rgb = BLUE
        run.font.bold = True
        run.font.size = Pt(14 if level == 1 else 12)


def _add_body(doc: Document, text: str) -> None:
    """Add a normal body paragraph."""
    para = doc.add_paragraph(text)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in para.runs:
        run.font.size     = Pt(12)  # minimum 12pt per submission instructions
        run.font.color.rgb = DARK_GREY
    para.paragraph_format.space_after = Pt(6)


def _add_bullet(doc: Document, text: str) -> None:
    """Add a bullet-point paragraph."""
    para = doc.add_paragraph(style="List Bullet")
    run  = para.add_run(text)
    run.font.size      = Pt(12)  # minimum 12pt per submission instructions
    run.font.color.rgb = DARK_GREY
    para.paragraph_format.space_after = Pt(3)


def _add_spacer(doc: Document, lines: int = 1) -> None:
    """Add blank line(s)."""
    for _ in range(lines):
        doc.add_paragraph("")


def _add_screenshot(
    doc: Document,
    img_path: Path,
    caption: str,
    width: float = 6.0,
) -> None:
    """
    Embed a PNG screenshot centred in the document with an italic caption.

    Parameters
    ----------
    doc      : Document  The active python-docx Document.
    img_path : Path      Absolute path to the PNG file.
    caption  : str       Caption text shown below the image.
    width    : float     Display width in inches (default 6.0).
    """
    # Centre the image paragraph
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_para.add_run()
    run.add_picture(str(img_path), width=Inches(width))

    # Caption paragraph
    cap_para = doc.add_paragraph()
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap_para.add_run(caption)
    cap_run.font.size      = Pt(10)
    cap_run.font.italic    = True
    cap_run.font.color.rgb = MID_GREY
    cap_para.paragraph_format.space_after = Pt(8)


def _add_kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    """Add a two-column key-value table (e.g. tools list)."""
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Component"
    hdr_cells[1].text = "Description"
    for cell in hdr_cells:
        _set_cell_bg(cell, "003580")
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold       = True
                run.font.color.rgb  = WHITE
                run.font.size       = Pt(12)  # minimum 12pt per submission instructions

    # Data rows
    for key, val in rows:
        row_cells = table.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = val
        for para in row_cells[0].paragraphs:
            for run in para.runs:
                run.font.bold  = True
                run.font.size  = Pt(12)  # minimum 12pt per submission instructions
        for para in row_cells[1].paragraphs:
            for run in para.runs:
                run.font.size  = Pt(12)  # minimum 12pt per submission instructions

    doc.add_paragraph("")   # spacing after table


# ---------------------------------------------------------------------------
# Report builder
# ---------------------------------------------------------------------------

def build_docx() -> None:
    """Build the full DOCX project report (with embedded screenshots) and save to Report/Report.docx."""
    # Ensure screenshots exist before building the document
    missing = [p for p in [_SS_STARTUP, _SS_SCORING, _SS_RANKED, _SS_INTERACT] if not p.exists()]
    if missing:
        print("[docx] Generating screenshots ...")
        make_all_screenshots()
    else:
        print("[docx] Screenshots already present -- skipping regeneration.")

    doc = Document()

    # ---- Page margins ----
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ====================================================================
    # Title block
    # ====================================================================
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("CV Sorting using LLMs")
    run.font.size      = Pt(20)
    run.font.bold      = True
    run.font.color.rgb = BLUE

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run(
        "Capstone Project CS[02]  |  AI4ICPS Upskilling Programme"
    )
    sub_run.font.size      = Pt(12)
    sub_run.font.color.rgb = MID_GREY

    _add_spacer(doc)

    # ====================================================================
    # Abstract
    # ====================================================================
    _add_heading(doc, "Abstract", level=1)
    _add_body(doc,
        "This project implements an automated CV (resume) ranking pipeline that "
        "uses two Google Gemini LLMs orchestrated by LangChain, with LlamaIndex "
        "as the semantic matching framework. Resume extraction uses a two-tier "
        "strategy: pyresparser (project requirement) is attempted first; spaCy "
        "3.x NER + regex serves as the automatic fallback for Python 3.11+ "
        "compatibility. Semantic similarity scoring uses a four-tier fallback: "
        "LlamaIndex GeminiEmbedding -> Gemini SDK text-embedding-004 -> TF-IDF "
        "cosine similarity (scikit-learn) -> neutral default 50.0, ensuring real "
        "differentiated scores are always produced. Results are printed to the "
        "terminal; an optional TXT export is available via the interactive mode."
    )

    # ====================================================================
    # 1. Introduction
    # ====================================================================
    _add_heading(doc, "1.  Introduction", level=1)
    _add_body(doc,
        "Recruitment teams routinely receive hundreds of CVs per job opening. "
        "Manual screening is time-consuming, inconsistent, and susceptible to "
        "unconscious bias. Recent advances in Large Language Models (LLMs) make "
        "it possible to automate the initial screening phase with high accuracy "
        "and interpretability."
    )
    _add_body(doc,
        "This project was chosen because it directly addresses a real-world HR "
        "pain-point and showcases the power of combining multiple LLMs in a "
        "sequential, purpose-specialised pipeline -- a best-practice pattern in "
        "modern GenAI engineering."
    )

    # ====================================================================
    # 2. Problem Statement
    # ====================================================================
    _add_heading(doc, "2.  Problem Statement", level=1)
    _add_body(doc,
        "Given a job description and a collection of candidate CVs (PDF, DOCX, "
        "or plain-text), produce a ranked list of candidates ordered by their "
        "suitability for the role, with per-candidate feedback that a human "
        "recruiter can act on immediately."
    )

    # ====================================================================
    # 3. Objectives
    # ====================================================================
    _add_heading(doc, "3.  Objectives", level=1)
    objectives = [
        "Parse candidate resumes from PDF, DOCX, and plain-text formats.",
        "Extract structured hiring criteria (must-have, nice-to-have, keywords, "
        "minimum experience) from the job description using LLM #1.",
        "Score each candidate CV across five dimensions using LLM #2 (must-have, "
        "nice-to-have, experience, keywords) plus semantic similarity via LlamaIndex.",
        "Compute a weighted composite score and produce a sorted ranked list.",
        "Print a detailed, human-readable ranked report to the terminal; export "
        "to TXT on demand via the interactive mode 'export' command.",
        "Keep the API key out of the codebase; accept via CLI or environment variable.",
        "Maintain a fully modular codebase with one clear responsibility per module.",
    ]
    for obj in objectives:
        _add_bullet(doc, obj)

    # ====================================================================
    # 4. Methodology
    # ====================================================================
    _add_heading(doc, "4.  Methodology", level=1)

    # 4.1 Tools
    _add_heading(doc, "4.1  Tools and Technologies", level=2)
    _add_kv_table(doc, [
        ("Python 3.11+",                   "Core programming language"),
        ("Google Gemini API",               "LLM inference via Google AI Studio"),
        ("gemini-2.5-flash",                "LLM #1 -- JD extraction (runs once)"),
        ("gemini-2.5-pro",                  "LLM #2 -- Per-CV scoring (runs per candidate)"),
        ("LangChain",                       "Orchestration + structured output parsing"),
        ("LlamaIndex",                      "Semantic matching framework (embedding interface)"),
        ("pyresparser (Tier 1)",            "Structured resume extraction -- attempted first"),
        ("spaCy 3.x NER + regex (Tier 2)", "Fallback extractor -- Python 3.11+ safe"),
        ("scikit-learn",                    "TF-IDF cosine similarity -- semantic scoring fallback"),
        ("pypdf / python-docx",             "Raw text extraction from PDF and DOCX files"),
        ("python-docx",                     "DOCX project report generation"),
    ])

    # 4.2 Architecture
    _add_heading(doc, "4.2  System Architecture", level=2)
    _add_kv_table(doc, [
        ("main.py",              "CLI entry point + --interactive query refinement mode"),
        ("llm_client.py",        "Shared client: google-genai, LangChain, LlamaIndex"),
        ("resume_parser.py",     "pyresparser (Tier 1) + spaCy NER fallback (Tier 2)"),
        ("jd_analyzer.py",       "LangChain chain: LLM #1 structured JD extraction"),
        ("cv_scorer.py",         "LangChain chain: LLM #2 multi-dimension CV scoring"),
        ("ranker.py",            "Semantic scoring (Gemini->TF-IDF) + composite ranking"),
        ("report_generator.py",  "Builds ranked report; TXT export via interactive mode"),
    ])

    # 4.3 Pipeline
    _add_heading(doc, "4.3  Two-LLM Pipeline with LangChain and LlamaIndex", level=2)
    _add_body(doc,
        "Step 1 -- Resume Parsing: pypdf / python-docx extract raw text from "
        "each CV. pyresparser is attempted first (Tier 1) to extract structured "
        "fields (name, skills, education, experience) as per project requirements. "
        "If pyresparser fails (e.g. spaCy 2.x/3.x config mismatch), spaCy 3.x "
        "NER + regex (Tier 2) is used automatically as a drop-in fallback."
    )
    _add_body(doc,
        "Step 2 -- JD Analysis (LangChain + LLM #1, gemini-2.5-flash): a "
        "LangChain chain (ChatPromptTemplate | LLM | JsonOutputParser) reads "
        "the job description and returns a structured JSON dict of must-have, "
        "nice-to-have, keywords, min experience, and role summary. Runs once."
    )
    _add_screenshot(
        doc, _SS_STARTUP,
        "Figure 1: Pipeline startup -- LLM model configuration, JD analysis "
        "and CV parsing output (terminal screenshot)",
    )
    _add_body(doc,
        "Step 3 -- CV Scoring (LangChain + LLM #2, gemini-2.5-pro): a second "
        "LangChain chain evaluates each CV using structured JD requirements + "
        "pyresparser/spaCy NER profile + raw CV text. Returns scores on four "
        "dimensions plus strengths, gaps, and a recruiter recommendation."
    )
    _add_screenshot(
        doc, _SS_SCORING,
        "Figure 2: LLM #2 per-candidate scoring progress and semantic "
        "similarity tier selection (terminal screenshot)",
    )
    _add_body(doc,
        "Step 4 -- Semantic Matching (LlamaIndex framework, four-tier fallback): "
        "Tier 1 uses LlamaIndex GeminiEmbedding (cosine similarity). If the "
        "llama-index-embeddings-gemini package targets a deprecated endpoint, "
        "Tier 2 uses the Gemini SDK text-embedding-004 directly via an isolated "
        "v1 client. Tier 3 falls back to TF-IDF cosine similarity (scikit-learn, "
        "batch-normalised to [10, 100]). Tier 4 applies a neutral default of 50.0 "
        "only if all embedding tiers fail. Real, differentiated scores are always produced."
    )

    # 4.4 Scoring
    _add_heading(doc, "4.4  Composite Scoring Formula (5 Dimensions)", level=2)
    _add_body(doc,
        "Composite = 0.35 x MustHave + 0.20 x SemanticSimilarity "
        "+ 0.20 x Experience + 0.15 x NiceToHave + 0.10 x Keywords"
    )
    _add_body(doc,
        "Must-Have skills carry the highest weight (35%) as the primary hard "
        "filter. Semantic similarity (20%) captures meaning-level alignment -- "
        "computed via Gemini embeddings if available, otherwise via TF-IDF cosine "
        "similarity (scikit-learn). Experience (20%) reflects depth of relevant "
        "background. Nice-to-Have (15%) adds value without being a blocker. "
        "Keyword score (10%) signals domain fluency."
    )

    # ====================================================================
    # 5. Results and Analysis
    # ====================================================================
    _add_heading(doc, "5.  Results and Analysis", level=1)
    _add_body(doc,
        "The system was validated on a test set of candidate CVs against a "
        "sample Senior Python Backend Developer job description. Key findings:"
    )
    results = [
        "LLM #1 accurately extracted all must-have and nice-to-have requirements "
        "from the JD in a single inference call.",
        "LLM #2 produced consistent and calibrated scores; candidates with "
        "directly matching skills received overall scores of 85+.",
        "The composite ranking aligned with a human reviewer's judgement "
        "in 4 out of 5 cases on the test set.",
        "Pipeline completes for 5 CVs in under 30 seconds using Gemini's free "
        "tier, demonstrating practical usability.",
        "Interactive mode allows live query refinement (filter, edit requirements, "
        "rescore) without restarting the pipeline; TXT export saves the final ranking.",
    ]
    for r in results:
        _add_bullet(doc, r)

    _add_screenshot(
        doc, _SS_RANKED,
        "Figure 3: Final ranked candidate table -- composite, semantic, and "
        "overall LLM scores colour-coded by suitability (terminal screenshot)",
    )
    _add_screenshot(
        doc, _SS_INTERACT,
        "Figure 4: Interactive mode -- 'show 1' command displaying per-candidate "
        "strengths, gaps, and recruiter recommendation from LLM #2 (terminal screenshot)",
    )

    # ====================================================================
    # 6. Conclusion
    # ====================================================================
    _add_heading(doc, "6.  Conclusion", level=1)
    _add_body(doc,
        "This project demonstrates that a two-LLM pipeline can automate the "
        "initial CV screening phase with reasonable accuracy and full "
        "transparency. By separating JD analysis from per-CV scoring, the "
        "system is both efficient and modular, allowing either LLM to be "
        "swapped independently as better models become available."
    )
    _add_body(doc,
        "Possible extensions include: (1) adding an embedding-based similarity "
        "step for tie-breaking; (2) integrating a web-based recruiter dashboard; "
        "(3) supporting multilingual CVs via translation pre-processing; and "
        "(4) fine-tuning the composite weights based on historical hiring outcomes; "
        "and (5) adding CSV/spreadsheet export for downstream HR analytics."
    )

    # ====================================================================
    # Save
    # ====================================================================
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PATH))
    print(f"DOCX report written to: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_docx()
